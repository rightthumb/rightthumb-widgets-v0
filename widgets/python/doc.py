#!/usr/bin/python3

# ## {R2D2919B742E} ##
# ###########################################################################
# What if magic existed?
# What if a place existed where your every thought and dream come to life.
# There is only one catch: it has to be written down.
# Such a place exists, it is called programming.
#    - Scott Taylor Reph, RightThumb.com
# ###########################################################################
# ## {C3P0D40fAe8B} ##


##################################################
import sys, time
##################################################
import _rightThumb._construct as __
appDBA=__.clearFocus(__name__,__file__);__.appReg=appDBA;
def focus(parentApp='',childApp='',reg=True):
	global appDBA;f=__.appName(appDBA,parentApp,childApp);
	if reg:__.appReg=f;
	return f
import _rightThumb._base3 as _
fieldSet=_.l.vars(focus(),__name__,__file__,appDBA)
_.load()
##################################################
_v = __.imp('_rightThumb._vars')
_str = __.imp('_rightThumb._string')
##################################################


def sw():
	pass
	#b)--> examples
	# _.switches.register( 'Input', '-i' )
	# _.switches.register( 'URL', '-u,-url,-urls', 'https://etc.ac/', isData='raw' )
	#e)--> examples
	_.switches.register( 'Files', '-f,-fi,-file,-files','file.txt', isData='name', description='Files', isRequired=False )

# __.setting('require-list',['Files,Plus','File,Has']) # todo
# __.setting('require-list',['Pipe','Files'])
__.setting('receipt-log')
__.setting('receipt-file')
__.setting('myFileLocations-skip-validation',False)
__.setting('require-pipe',False)
__.setting('require-pipe||file',False)
__.setting('pre-error',False)
__.setting('switch-raw',[])



_.appInfo[focus()] = {
	# 'app': '8facG-jo0Cxk',
	'file': 'thisApp.py',
	'liveAppName': __.thisApp( __file__ ),
	'description': 'Changes the world',
		# _.ail(1,'subject')+
		# _.aib('one')+
	'categories': [
						'DEFAULT',
				],
	'usage': [
						# 'epy another',
						# 'e nmap',
						# '',
	],
	'relatedapps': [
						# 'p another -file file.txt',
						# '',
	],
	'prerequisite': [
						# 'p another -file file.txt',
						# '',
	],
	'examples': [
						_.hp('p thisApp -file file.txt'),
						_.linePrint(label='simple',p=0),
						'',
	],
	'columns': [
					# { 'name': 'name', 'abbreviation': 'n' },
					# { 'name': '{1}', 'abbreviation': '{0}', 'sort': '{2}' },
	],
	'aliases': [
					# 'this',
					# 'app',
	],
	'notes': [
					# {},
	],
}

_.appData[focus()] = {
		'start': __.startTime,
		'uuid': '',
		'audit': [],
		'pipe': False,
		'data': {
					'field': {'sent': [], 'received': [] }, # { 'label': '', 'context': [],  }
					'table': {'sent': [], 'received': [] },
		},
	}


def triggers():
	_.switches.trigger( 'Files', _.myFileLocations, vs=True )
	_.switches.trigger( 'Ago', _.timeAgo )
	_.switches.trigger( 'Folder', _.myFolderLocations )
	_.switches.trigger( 'URL', _.urlTrigger )
	_.switches.trigger( 'Duration', _.timeFuture )

def _local_(do): exec(do)

_.l.conf('clean-pipe',True)
_.l.sw.register( triggers, sw )

########################################################################################
#b)--> examples
#d)--> code hints to quickly get started
	#n)--> inline examples
		# any(ele in 'scott5' for ele in list('0123456789'))
		# if _.switches.isActive('Test'): test(); return None;
		# result=[]; result=[ _.pr(line) for i, line, bi in _.numerate( _.isData(r=0) )]
		# bk=[];[  bk.append(rec['backup']) for rec in backupLog if path == rec['file']]; bk=bk[-1];
		# a=(1 if True else 0) <--# 
		#!)--> m=[[row[i] for row in matrix] for i in range(4)]

	#n)--> python globals
		# globals()['var']
		# for k in globals(): print(k, eval(k) )

	#n)--> webpage from url
		# for subject in _.caseUnspecific( line, needle ): line = line.replace( subject, _.colorThis( subject, 'green', p=0 ) )

	#n)--> webpage from url
		# requests=__.imp('requests.post')
		#!)--> data=str(requests.post(url,data={}).content,'iso-8859-1')

	#n)--> import and backup example
		# _bk = _.regImp( __.appReg, 'fileBackup' ); _bk.switch( 'Silent' ); _bk.switch( 'isRunOnce' ); _bk.switch( 'Flag', 'APP' ); _bk.switch( 'DoNotSchedule' )
		# _bk.switch( 'Input', path ); bkfi = _bk.action();
	
	#n)--> inline
		# for rel in [ subject for subject in _.isData(r=0) if _.showLine(subject) ]: print(rel)

	#n)--> banner
		# banner=_.Banner(app); goss=banner.goss;
#e)--> examples
########################################################################################
#n)--> start

# pip install textract
# pip install docx2txt
# pip install PyPDF2
# pip install pdfminer.six


# try: import textract
# except: pass

# # try: import docx2txt
# # except: pass

# try: import win32com.client
# except: pass

# # try: from PyPDF2 import PdfReader
# # except: pass

# # try: import PyPDF2
# # except: pass

# try: from pathlib import Path
# except: pass

# try: import docx
# except: pass

# import os
# # import textract

import re
import docx
from pathlib import Path
import textract

def extract_headers_footers(doc):
	headers = []
	footers = []

	for section in doc.sections:
		for paragraph in section.header.paragraphs:
			headers.append(paragraph.text)
		for paragraph in section.footer.paragraphs:
			footers.append(paragraph.text)

	return "\n".join(headers), "\n".join(footers)

def extract_phone_numbers2(text):
	phone_regex = re.compile(r'\b\d{3}[-\.\s]??\d{3}[-\.\s]??\d{4}\b')
	return phone_regex.findall(text)

def extract_phone_numbers(text):
	phone_regex = re.compile(r'\(?\b\d{3}\)?[-\.\s]??\d{3}[-\.\s]??\d{4}\b')
	return phone_regex.findall(text)


def extract_emails(text):
	email_regex = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
	return email_regex.findall(text)

def extract_addresses(text):
	address_regex = re.compile(r'\d+\s[A-Za-z0-9]+\s[A-Za-z0-9\s,]*[A-Za-z]\s\d{5}(-\d{4})?')
	return address_regex.findall(text)

def getContents2(file_path):
	path = Path(file_path)
	text = ""

	try:
		if path.suffix == ".docx":
			doc = docx.Document(file_path)
			full_text = []

			for paragraph in doc.paragraphs:
				full_text.append(paragraph.text)

			headers, footers = extract_headers_footers(doc)
			# print(headers, footers)
			text = "\n".join([headers, "\n".join(full_text), footers])
		else:
			text = textract.process(file_path).decode("utf-8")
	except Exception as e:
		print(f"Error processing {file_path}: {e}")
	print(text)

	phone_numbers = extract_phone_numbers(text)
	emails = extract_emails(text)
	addresses = extract_addresses(text)

	return text, phone_numbers, emails, addresses

# file_path = "your_file_path_here.docx"
# text, phone_numbers, emails, addresses = getContents2(file_path)
# print("Phone numbers:", phone_numbers)
# print("Emails:", emails)
# print("Addresses:", addresses)






def action():
	for path in _.isData(r=0):
		# doc=getContents2(path)
		# _.pr(doc)
		file_path = path
		text, phone_numbers, emails, addresses = getContents2(file_path)
		print("Phone numbers:", phone_numbers)
		print("Emails:", emails)
		print("Addresses:", addresses)




##################################################
#b)--> examples
# banner=_.Banner(dependencies)
# goss=banner.goss
# goss('-\t this app will sherlock tf out of any python app or python module')
#e)--> examples
##################################################

########################################################################################
if __name__ == '__main__':
	#b)--> examples

	# banner.pr()
	# if len(_.switches.all())==0: banner.gossip()
	
	#e)--> examples
	action()
	_.isExit(__file__)